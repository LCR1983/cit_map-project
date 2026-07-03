import docx

file_path = "郷土料理一覧_504件.docx"
doc = docx.Document(file_path)

print(f"Total tables: {len(doc.tables)}")

for i, table in enumerate(doc.tables):
    print(f"Table {i} has {len(table.rows)} rows and {len(table.columns)} columns.")
    if len(table.rows) > 0:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        print(f"  Headers: {headers}")
    # Print first few rows
    for r in range(1, min(5, len(table.rows))):
        row_data = [cell.text.strip() for cell in table.rows[r].cells]
        print(f"  Row {r}: {row_data}")
